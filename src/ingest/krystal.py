"""
Krystal DOCX preseed: parses the Krystal OP glossary + style guide.

Sources:
  sources/czech/krystal/Teologická Suma úzus - verze 4.docx

Produces:
  - glossary_term rows (one per Latin lemma)
  - glossary_sense rows (one per sense; multi-sense terms get N rows)
  - sense_rendering(cs) rows from DOCX Czech renderings
  - sense_rendering(sk) rows using Czech as placeholder (status='proposed')

Sense detection:
  Pattern 1 — comma-separated senses: "term1 (context1), term2 (context2)"
  Pattern 2 — colon-in-paren:          "primary (context: alternative)"
  Single-sense note:                    "term (też: alt)"  →  treated as single-sense

Run (shows extraction counts for review, then inserts after approval):
  uv run python -m ingest.krystal
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import docx

from common.db import get_conn, source_id

ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "sources" / "czech" / "krystal" / "Teologická Suma úzus - verze 4.docx"

# Paragraph index where the glossary begins (heading "SLOVNÍČEK TERMÍNŮ")
_GLOSSARY_HEADER = "SLOVNÍČEK TERMÍNŮ"

# Terms the spec identifies as known multi-sense (non-exhaustive; from m1_resolution.md).
# Any of these that appear in the DOCX must parse as multi-sense; validation warns if not.
# Terms absent from the DOCX entirely (bonum, actus, potentia, intentio) are gap terms
# resolved by the resolver — they are expected to be missing here.
_KNOWN_MULTISENSE = frozenset({
    "concupiscentia", "gratia", "fides", "intellectus", "providentia",
    "ratio", "passio", "forma", "virtus", "bonum", "actus", "potentia",
    "species", "intentio", "sensus",
})

# Words in a parenthetical that indicate "also/alternatively" — NOT a distinct sense
_ALSO_MARKERS = re.compile(
    r"^(?:taky|také|též|ale i|někdy(?:\s+i|\s+taky|\s+také|\s+též)?|popř\.|případně|dle kontextu|nebo)\b",
    re.IGNORECASE,
)

# Trailing "also" qualifiers to strip from parsed context labels.
# Czech "též" = t + é (é) + ž (ž); use escapes to avoid source-encoding confusion.
_TRAILING_ALSO = re.compile(
    " +(?:také|taky|též|i) *$", re.IGNORECASE
)

# ── Data model ────────────────────────────────────────────────────────────────


@dataclass
class Sense:
    context_label: str | None  # NULL = primary / only sense
    cs_rendering: str          # Czech rendering (also used as sk placeholder)


@dataclass
class GlossaryEntry:
    latin_lemma: str
    is_multiword: bool
    senses: list[Sense]
    raw_text: str = field(repr=False)


# ── Separator splitting ───────────────────────────────────────────────────────

_DASH_RE = re.compile(r"\s[–-]\s")


def _split_latin_czech(line: str) -> tuple[str, str] | None:
    """Split 'latin – czech' at the first em-dash or hyphen separator.
    Returns (latin, czech) or None if no separator found.
    """
    m = _DASH_RE.search(line)
    if m is None:
        return None
    latin = line[: m.start()].strip()
    czech = line[m.end() :].strip()
    return latin, czech


def _clean_latin(raw: str) -> tuple[str, bool]:
    """Return (lemma, is_multiword) from the raw Latin side.

    Strips parenthetical type-annotations like 'species (impressa, intelligibilis)'.
    Treats multi-word lemmas (contains space) as is_multiword=True.
    """
    # Remove trailing parenthetical variants: 'species (impressa, intelligibilis)'
    lemma = re.sub(r"\s*\([^)]*\)\s*$", "", raw).strip()
    is_multiword = " " in lemma
    return lemma, is_multiword


# ── Multi-sense detection ─────────────────────────────────────────────────────


def _parse_senses(czech_content: str) -> list[Sense]:
    """Parse Czech content into a list of Senses.

    Handles three patterns:
      (a) Comma outside parens → Pattern 1: multiple labeled senses
          "žádostivost (label1), dychtění (label2)"
      (b) Colon inside paren with domain context → Pattern 2
          "primary (context: alternative)"
      (c) Everything else → single sense, full text as cs_rendering
    """
    # --- Pattern 1: comma OUTSIDE parentheses separating distinct senses ------
    # Split on commas that are not inside parentheses
    outside_comma_parts = _split_outside_parens(czech_content, ",")
    if len(outside_comma_parts) >= 2:
        # Each part should be "term (context)" or just "term"
        senses = []
        for part in outside_comma_parts:
            part = part.strip()
            term, label = _extract_term_and_label(part)
            if term:
                senses.append(Sense(context_label=label or None, cs_rendering=term))
        # Only treat as multi-sense when at least one sense has a context label.
        # Comma-separated synonyms (all NULL labels) e.g. "nevědomost, neznalost"
        # are the same concept and should be collapsed to a single-sense entry.
        if len(senses) >= 2 and any(s.context_label is not None for s in senses):
            return senses

    # --- Pattern 2: colon inside parenthetical (domain: alternative_term) -----
    colon_match = re.search(r"\(([^:)]+):\s*([^)]+)\)", czech_content)
    if colon_match:
        paren_context = colon_match.group(1).strip()
        paren_alternative = colon_match.group(2).strip()
        # Skip if the context text starts with an "also" marker
        if not _ALSO_MARKERS.match(paren_context):
            # Primary rendering: everything before the parenthetical
            primary = czech_content[: colon_match.start()].strip().rstrip(",").strip()
            # Clean trailing "also" qualifiers from context label
            # e.g. "v případě ctnosti též" → "v případě ctnosti"
            clean_context = _TRAILING_ALSO.sub("", paren_context).strip()
            if primary and paren_alternative:
                return [
                    Sense(context_label=None, cs_rendering=primary),
                    Sense(context_label=clean_context, cs_rendering=paren_alternative),
                ]

    # --- Single sense (including parenthetical clarifications) ----------------
    # If there are comma-separated synonyms with no context (fell through above),
    # use the first synonym as the canonical rendering.
    outside_parts = _split_outside_parens(czech_content, ",")
    if len(outside_parts) > 1:
        czech_content = outside_parts[0].strip()

    # Strip trailing parenthetical notes for the clean rendering
    clean = re.sub(r"\s*\([^)]*\)\s*$", "", czech_content).strip()
    cs = clean if clean else czech_content.strip()
    return [Sense(context_label=None, cs_rendering=cs)]


def _split_outside_parens(text: str, delimiter: str) -> list[str]:
    """Split `text` on `delimiter` only when not inside parentheses."""
    parts = []
    depth = 0
    current = []
    for ch in text:
        if ch == "(":
            depth += 1
            current.append(ch)
        elif ch == ")":
            depth -= 1
            current.append(ch)
        elif ch == delimiter and depth == 0:
            parts.append("".join(current))
            current = []
        else:
            current.append(ch)
    if current:
        parts.append("".join(current))
    return parts


def _extract_term_and_label(part: str) -> tuple[str, str | None]:
    """From 'term (label)' extract (term, label). From 'term' return (term, None)."""
    m = re.match(r"^(.*?)\s*\(([^)]+)\)\s*$", part.strip())
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return part.strip(), None


# ── Glossary parsing ──────────────────────────────────────────────────────────


def parse_glossary(doc: docx.Document) -> list[GlossaryEntry]:
    """Parse all glossary entries from the SLOVNÍČEK TERMÍNŮ section."""
    # Find the start of the glossary section
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if _GLOSSARY_HEADER in p.text:
            start_idx = i + 1
            break
    if start_idx is None:
        raise RuntimeError(f"FAIL: '{_GLOSSARY_HEADER}' heading not found in DOCX")

    entries: list[GlossaryEntry] = []
    seen_lemmas: dict[str, int] = {}  # lemma → index in entries (for dedup)

    for p in doc.paragraphs[start_idx:]:
        text = p.text.strip()
        if not text:
            continue

        result = _split_latin_czech(text)
        if result is None:
            continue  # skip non-entry lines

        latin_raw, czech_content = result
        if not latin_raw or not czech_content:
            continue

        # Handle comma-separated Latin lemmas: "veracitas, veritas – pravdivost"
        # Use outside-parens split to avoid splitting "species (impressa, intelligibilis)"
        latin_lemmas_raw = [part.strip() for part in _split_outside_parens(latin_raw, ",")]

        for lat_raw in latin_lemmas_raw:
            lemma, is_multiword = _clean_latin(lat_raw)
            if not lemma:
                continue

            senses = _parse_senses(czech_content)
            entry = GlossaryEntry(
                latin_lemma=lemma,
                is_multiword=is_multiword,
                senses=senses,
                raw_text=text,
            )

            if lemma in seen_lemmas:
                # Later entry supersedes earlier one (e.g. "veritas" appears twice)
                print(f"  NOTE: duplicate lemma {lemma!r} — later entry supersedes", flush=True)
                entries[seen_lemmas[lemma]] = entry
            else:
                seen_lemmas[lemma] = len(entries)
                entries.append(entry)

    return entries


# ── Review output ─────────────────────────────────────────────────────────────


def print_review(entries: list[GlossaryEntry]) -> None:
    """Print extraction summary for human review before DB insert.

    Also validates that known multi-sense terms present in the glossary were
    actually parsed as multi-sense — warns if any were silently collapsed.
    """
    multi = [e for e in entries if len(e.senses) > 1]
    single = [e for e in entries if len(e.senses) == 1]
    multiword = [e for e in entries if e.is_multiword]
    total_senses = sum(len(e.senses) for e in entries)
    lemma_to_entry = {e.latin_lemma: e for e in entries}

    print(f"Terms:          {len(entries)}")
    print(f"Total senses:   {total_senses}")
    print(f"Single-sense:   {len(single)}")
    print(f"Multi-sense:    {len(multi)}")
    print(f"Multi-word:     {len(multiword)}")
    print()

    print("Multi-sense terms:")
    for e in multi:
        print(f"  {e.latin_lemma}:")
        for s in e.senses:
            label = f"[{s.context_label}]" if s.context_label else "[primary]"
            print(f"    {label} → {s.cs_rendering!r}")
    print()

    print("Multi-word terms:")
    for e in multiword:
        print(f"  {e.latin_lemma!r}")
    print()

    # Validate: known multi-sense terms that appear in the glossary should parse as multi-sense
    in_glossary = _KNOWN_MULTISENSE & lemma_to_entry.keys()
    parsed_single = [t for t in in_glossary if len(lemma_to_entry[t].senses) == 1]
    absent = _KNOWN_MULTISENSE - lemma_to_entry.keys()
    if parsed_single:
        print(f"WARN: {len(parsed_single)} known multi-sense term(s) parsed as single-sense "
              f"— check DOCX format or _parse_senses logic:")
        for t in sorted(parsed_single):
            print(f"  {t!r}: {lemma_to_entry[t].raw_text!r}")
    if absent:
        print(f"INFO: {len(absent)} known multi-sense term(s) absent from DOCX "
              f"(gap terms — resolver will handle):")
        print(f"  {', '.join(sorted(absent))}")


# ── DB insertion ──────────────────────────────────────────────────────────────


def insert_glossary(conn, entries: list[GlossaryEntry], src_krystal: int) -> dict:
    """Insert glossary_term, glossary_sense, sense_rendering rows.

    Returns counts: {terms, senses, renderings}.
    Idempotent: truncates existing rows before inserting.
    """
    cur = conn.cursor()

    # Wipe existing data (idempotency)
    cur.execute("DELETE FROM sense_rendering")
    cur.execute("DELETE FROM glossary_sense")
    cur.execute("DELETE FROM glossary_term")

    term_count = sense_count = rendering_count = 0

    for entry in entries:
        cur.execute(
            """
            INSERT INTO glossary_term (latin_lemma, is_multiword)
            VALUES (%s, %s)
            RETURNING term_id
            """,
            (entry.latin_lemma, entry.is_multiword),
        )
        term_id = cur.fetchone()[0]
        term_count += 1

        for sense in entry.senses:
            # status='approved': Krystal is the terminological authority for the SENSE
            # DEFINITION (what the Latin term means). The SK rendering being a Czech
            # placeholder is a separate quality concern handled in the M3 review pass.
            cur.execute(
                """
                INSERT INTO glossary_sense (term_id, context_label, status)
                VALUES (%s, %s, 'approved')
                RETURNING sense_id
                """,
                (term_id, sense.context_label),
            )
            sense_id = cur.fetchone()[0]
            sense_count += 1

            # cs rendering — from DOCX.
            # lemma = cs_rendering directly (Krystal renderings are dictionary forms).
            # TODO(M2): run MorphoDiTa on cs_rendering to get the true lemma form,
            # in case any rendering is not in nominative singular (e.g. compound forms).
            cur.execute(
                """
                INSERT INTO sense_rendering (sense_id, lang, lemma, content, source_id)
                VALUES (%s, 'cs', %s, %s, %s)
                """,
                (sense_id, sense.cs_rendering, sense.cs_rendering, src_krystal),
            )
            rendering_count += 1

            # sk rendering — Czech content as placeholder pending Slovak theologian review.
            # source_id=krystal is intentional for M1 (resolver reads sk rows regardless of
            # source). TODO(M4): replace with actual Slovak content and source_id=model
            # before the translation prompt injects this as a hard constraint. Until then
            # the translator will receive Czech terms as Slovak constraints — reviewable.
            cur.execute(
                """
                INSERT INTO sense_rendering (sense_id, lang, content, source_id)
                VALUES (%s, 'sk', %s, %s)
                """,
                (sense_id, sense.cs_rendering, src_krystal),
            )
            rendering_count += 1

            # TODO(M2): add en cue rows once a Latin→English term mapping is available.
            # EN cues serve as disambiguation signals for multi-sense resolution
            # (resolver Step 7). For M1, Czech signals from Bahounek text are sufficient
            # and the English signal path will simply produce no evidence (not an error).

    cur.close()
    return {"terms": term_count, "senses": sense_count, "renderings": rendering_count}


# ── Entry point ───────────────────────────────────────────────────────────────


def run(*, skip_confirm: bool = False) -> None:
    """Parse DOCX, print review summary, then insert after confirmation."""
    if not DOCX_PATH.exists():
        raise RuntimeError(f"FAIL: Krystal DOCX not found at {DOCX_PATH}")

    print(f"Parsing {DOCX_PATH.name} ...")
    doc = docx.Document(DOCX_PATH)
    entries = parse_glossary(doc)

    print()
    print_review(entries)

    if not skip_confirm:
        print("\nReview the above. Type 'yes' to insert into DB, anything else to abort:")
        answer = input("> ").strip().lower()
        if answer != "yes":
            print("Aborted.")
            return

    # Insert into DB
    with get_conn() as conn:
        src = source_id(conn, "krystal")
        counts = insert_glossary(conn, entries, src)

    print(
        f"Inserted: {counts['terms']} terms, "
        f"{counts['senses']} senses, "
        f"{counts['renderings']} renderings"
    )


if __name__ == "__main__":
    try:
        run()
    except RuntimeError as exc:
        print(f"\n{exc}", file=sys.stderr)
        sys.exit(1)
